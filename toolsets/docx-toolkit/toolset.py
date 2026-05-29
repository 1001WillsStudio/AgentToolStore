"""
docx‑toolkit — Create, edit, extract, and analyze Word (.docx) documents.
============================================================================

Read text, inspect structure, extract tables, manage comments and tracked
changes — without needing to write throw‑away scripts for every document task.
"""

import io
from pathlib import Path

try:
    from toolstore.toolset import tool
except ImportError:
    def tool(fn):
        return fn


# ── docx_read ──────────────────────────────────────────────────────────

@tool
def docx_read(*, filepath: str) -> dict:
    """Extract all text from a Word document, preserving paragraph structure.

    Args:
        filepath: Path to the .docx file.

    Returns:
        dict with:
          text       — full document text (paragraphs separated by newlines)
          paragraphs — number of paragraphs
          sections   — number of sections
    """
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed — run: pip install python-docx"}

    p = Path(filepath).expanduser().resolve()
    if not p.exists():
        return {"error": f"File not found: {p}"}

    try:
        doc = Document(str(p))
    except Exception as exc:
        return {"error": f"Cannot open document: {exc}"}

    paragraphs = [para.text for para in doc.paragraphs]
    return {
        "text": "\n".join(paragraphs),
        "paragraphs": len(paragraphs),
        "sections": len(doc.sections),
    }


# ── docx_info ──────────────────────────────────────────────────────────

@tool
def docx_info(*, filepath: str) -> dict:
    """Inspect document structure: paragraphs, tables, sections, styles, images.

    Args:
        filepath: Path to the .docx file.

    Returns:
        dict with: paragraphs, tables, sections, styles_count, images,
                   core_properties (title, author, etc. if available)
    """
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed — run: pip install python-docx"}

    p = Path(filepath).expanduser().resolve()
    if not p.exists():
        return {"error": f"File not found: {p}"}

    try:
        doc = Document(str(p))
    except Exception as exc:
        return {"error": f"Cannot open document: {exc}"}

    # Count inline images
    image_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                image_count += 1
            if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
                image_count += 1

    props = {}
    try:
        cp = doc.core_properties
        props = {
            "title": cp.title or None,
            "author": cp.author or None,
            "created": str(cp.created) if cp.created else None,
            "modified": str(cp.modified) if cp.modified else None,
            "last_modified_by": cp.last_modified_by or None,
        }
    except Exception:
        pass

    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "styles_count": len(doc.styles) if doc.styles else 0,
        "images": image_count,
        "core_properties": props,
        "filename": p.name,
    }


# ── docx_extract_tables ────────────────────────────────────────────────

@tool
def docx_extract_tables(*, filepath: str, table_index: int = 0) -> dict:
    """Extract tables from a Word document as structured data.

    Args:
        filepath:    Path to the .docx file.
        table_index: Which table to extract (0 = first, -1 = all).

    Returns:
        dict with:
          tables — list of table dicts: {index, rows, columns, data: [[...]]}
          count  — total tables in document
    """
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed — run: pip install python-docx"}

    p = Path(filepath).expanduser().resolve()
    if not p.exists():
        return {"error": f"File not found: {p}"}

    try:
        doc = Document(str(p))
    except Exception as exc:
        return {"error": f"Cannot open document: {exc}"}

    if not doc.tables:
        return {"tables": [], "count": 0}

    indices = range(len(doc.tables)) if table_index == -1 else [table_index]
    tables = []
    for idx in indices:
        if idx < 0 or idx >= len(doc.tables):
            continue
        tbl = doc.tables[idx]
        rows = []
        for row in tbl.rows:
            cells = [cell.text for cell in row.cells]
            rows.append(cells)
        tables.append({
            "index": idx,
            "rows": len(rows),
            "columns": len(rows[0]) if rows else 0,
            "data": rows,
        })

    return {"tables": tables, "count": len(doc.tables)}


# ── docx_create ────────────────────────────────────────────────────────

@tool
def docx_create(*, filepath: str, title: str = "",
                content: list = None, author: str = "") -> dict:
    """Create a new Word document from structured content.

    Args:
        filepath: Where to write the new .docx file.
        title:    Document title (appears as a heading).
        content:  List of strings — each becomes a paragraph.
        author:   Document author metadata.

    Returns:
        dict with "written" (absolute path) or "error".
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        return {"error": "python-docx not installed — run: pip install python-docx"}

    p = Path(filepath).expanduser().resolve()
    p.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Core properties
    if author:
        try:
            doc.core_properties.author = author
        except Exception:
            pass

    # Title
    if title:
        h = doc.add_heading(title, level=1)

    # Content paragraphs
    for item in (content or []):
        if isinstance(item, str):
            if item.startswith("# "):
                doc.add_heading(item[2:], level=2)
            elif item.startswith("## "):
                doc.add_heading(item[3:], level=3)
            elif item.startswith("- "):
                doc.add_paragraph(item[2:], style="List Bullet")
            else:
                doc.add_paragraph(item)
        elif isinstance(item, dict):
            # Table from dict format
            rows_data = item.get("rows", [])
            cols_data = item.get("columns", [])
            if rows_data and cols_data:
                table = doc.add_table(rows=len(rows_data) + 1, cols=len(cols_data))
                table.style = "Light Grid Accent 1"
                for ci, col_name in enumerate(cols_data):
                    table.cell(0, ci).text = col_name
                for ri, row in enumerate(rows_data):
                    for ci, col_name in enumerate(cols_data):
                        table.cell(ri + 1, ci).text = str(row.get(col_name, ""))

    doc.save(str(p))
    return {"written": str(p), "paragraphs": len(doc.paragraphs)}
